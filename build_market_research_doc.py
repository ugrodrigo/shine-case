"""Build the source-linked Shine external research note as a Word document."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from build_final_doc import (
    BLACK,
    LIGHT_BLUE,
    LIGHT_GREY,
    MID_GREY,
    NAVY,
    RED,
    TEAL,
    WHITE,
    add_body,
    add_bullet,
    add_cell_text,
    add_heading,
    add_hyperlink,
    add_run,
    configure_document,
    prevent_row_split,
    set_cell_margins,
    set_cell_shading,
    set_repeat_table_header,
    set_table_borders,
)


ROOT = Path(__file__).resolve().parent
OUTPUT_FILE = ROOT / "Shine_External_Market_Research_and_Revised_Strategy.docx"


def add_title(document):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "SHINE • FRANCE", bold=True, color=TEAL, size=10)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    add_run(p, "External market research & revised strategy", bold=True, color=NAVY, size=23)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    add_run(
        p,
        "Pressure-testing the dataset-led BTP and initial-plan findings against Shine's current product, competition, market conditions, and regulation",
        color=MID_GREY,
        size=10.5,
    )
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p, "Research date: 30 August 2026", bold=True, color=BLACK, size=8.5)


def add_source_line(document, links):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Sources: ", bold=True, color=TEAL, size=7.4)
    for index, (label, url) in enumerate(links):
        if index:
            add_run(p, " • ", color=MID_GREY, size=7.4)
        add_hyperlink(p, label, url, size=7.4)
    return p


def add_decision_box(document):
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(18.3)
    cell = table.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=160, start=190, bottom=160, end=190)
    add_cell_text(cell, "Revised position", bold=True, color=WHITE, size=9.5)
    add_cell_text(
        cell,
        "BTP is Shine's strongest demonstrated vertical—not yet a proven untapped market. Run a bounded BTP-fit workflow experiment, retain behavior-based revenue protection as the second initiative, and use electronic invoicing as a cross-persona entry point only when causal tests show incremental value.",
        bold=True,
        color=WHITE,
        size=11.2,
        space_after=0,
    )
    set_table_borders(table, color=NAVY, size="1")


def add_evidence_table(document):
    add_heading(document, "What the supplied dataset establishes", level=1)
    rows = [
        ("BTP", "26.74% of comparable revenue", "N=1,642; 1.48× overall average revenue/company"),
        ("Plus initial plan", "28.56% of comparable plan revenue", "2.02× overall average revenue/company"),
        ("Start initial plan", "44.64% of comparable plan revenue", "Scale-led; approximately average revenue/company"),
        ("Business initial plan", "6.95% of comparable plan revenue", "N=196; 3.22× average, but small and fragile"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Cm(4.0), Cm(5.7), Cm(8.6)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    for idx, header in enumerate(("Segment", "Observed scale", "Interpretation")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=80, start=100, bottom=80, end=100)
        add_cell_text(cell, header, bold=True, color=WHITE, size=8.2, space_after=0)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(values):
            set_cell_shading(cells[idx], WHITE if row_index % 2 == 0 else LIGHT_GREY)
            set_cell_margins(cells[idx], top=70, start=95, bottom=70, end=95)
            add_cell_text(cells[idx], value, bold=(idx == 0), color=NAVY if idx == 0 else BLACK, size=8.0, space_after=0)
    set_table_borders(table, color="D8DEE3", size="3")
    add_body(
        document,
        "Interpretation boundary: these are associations inside Shine's extracted population. They do not establish total addressable market, current plan, acquisition cost, contribution margin, competitive switching, or causal product effects.",
        size=8.2,
        color=RED,
        space_after=5,
    )


def build_page_one(document):
    add_title(document)
    add_decision_box(document)
    add_evidence_table(document)
    add_heading(document, "Why the wording changed", level=1)
    add_bullet(document, "Shine already markets a dedicated BTP proposition, so strong BTP results may partly reflect an existing strategy or selection effect.")
    add_bullet(document, "The French construction base is large and small-business-heavy, supporting product fit but not proving unused headroom.")
    add_bullet(document, "Sector contraction, working-capital needs, and advanced BTP invoicing requirements narrow the ideal customer profile.")
    add_bullet(document, "Electronic invoicing creates a broader cross-persona trigger that the supplied dataset cannot size.")


def build_page_two(document):
    add_heading(document, "1. Shine's business model and strategic direction", level=1, space_before=0)
    add_body(
        document,
        "Shine is an ACPR-authorised payment institution rather than a traditional bank. It combines a French IBAN, cards and transfers with invoicing, receipt management, accounting integrations, expense controls, and administrative tools. It cannot directly provide an authorised overdraft; some financing depends on partners.",
        size=9.0,
    )
    add_body(
        document,
        "The visible model combines tier subscriptions and usage fees. Current headline monthly prices excluding tax are Free at EUR 0, Start at EUR 9, Plus at EUR 20, and Business at EUR 60, subject to billing terms and annual discounts. Business is a team product with ten premium physical cards, unlimited virtual cards, and higher quotas.",
        size=9.0,
    )
    add_body(
        document,
        "Following the Ageras acquisition, the direction is broader: a European small-business platform integrating payments, invoicing, accounting, payroll, and tax administration. That favors workflow depth and recurring administrative engagement over competing only as a current account.",
        size=9.0,
    )
    add_source_line(document, [
        ("Shine pricing", "https://www.shine.fr/tarifs/"),
        ("Payment institution", "https://help.shine.fr/fr/articles/1175387-quelles-sont-les-differences-entre-shine-et-une-banque-dite-traditionnelle"),
        ("Overdraft policy", "https://help.shine.fr/fr/articles/1179200-compte-a-decouvert"),
        ("Ageras", "https://help.shine.fr/fr/articles/10289941-shine-rejoint-ageras"),
    ])

    add_heading(document, "2. Competitive pressure", level=1)
    competitors = [
        ("Shine", "Free; EUR 9; EUR 20; EUR 60", "French service, invoicing, administration, BTP proposition"),
        ("Qonto", "EUR 9; EUR 19; EUR 39", "Accounting integration, cash-flow tools, expense management"),
        ("Indy", "Core account free; Premium from EUR 12", "Free account integrated with accounting and declarations"),
        ("Revolut Business", "EUR 10; EUR 35; EUR 125", "International payments, FX, team controls"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    for idx, width in enumerate((Cm(4.1), Cm(5.4), Cm(8.8))):
        table.columns[idx].width = width
    for idx, header in enumerate(("Provider", "Indicative pricing", "Competitive angle")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=75, start=95, bottom=75, end=95)
        add_cell_text(cell, header, bold=True, color=WHITE, size=8.1, space_after=0)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(competitors):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(values):
            set_cell_shading(cells[idx], WHITE if row_index % 2 == 0 else LIGHT_GREY)
            set_cell_margins(cells[idx], top=65, start=90, bottom=65, end=90)
            add_cell_text(cells[idx], value, bold=(idx == 0), color=NAVY if idx == 0 else BLACK, size=7.9, space_after=0)
    set_table_borders(table, color="D8DEE3", size="3")
    add_body(document, "Prices are not perfectly comparable: cards, quotas, billing periods, support, and included software differ. The strategic signal is that basic banking and e-invoicing are crowded; differentiation must come from workflow depth, service, or vertical relevance.", size=8.0, color=MID_GREY)
    add_source_line(document, [
        ("Qonto", "https://qonto.com/fr/pricing"),
        ("Indy", "https://www.indy.fr/compte-pro/"),
        ("Revolut", "https://www.revolut.com/fr-FR/business/business-account-plans/"),
    ])


def build_page_three(document):
    add_heading(document, "3. BTP: substantial opportunity, bounded fit", level=1, space_before=0)
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(9.15)
    table.columns[1].width = Cm(9.15)
    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, top=125, start=145, bottom=125, end=145)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(left, LIGHT_BLUE)
    set_cell_shading(right, LIGHT_GREY)
    add_cell_text(left, "Why it can scale", bold=True, color=TEAL, size=10)
    for item in (
        "650,207 active construction legal entities in France in 2024.",
        "Approximately 95% of building businesses are artisan-sized.",
        "Strong fit with mobile banking, invoices, collections, expense cards, subaccounts, and accounting.",
        "BTP combines scale and above-average Shine revenue in the dataset.",
    ):
        add_cell_text(left, "• " + item, size=8.4)
    add_cell_text(right, "What caps it", bold=True, color=RED, size=10)
    for item in (
        "Artisan activity fell 3.8% in 2025 and remained down 1.5% YoY in Q1 2026.",
        "Shine cannot directly provide an authorised overdraft.",
        "Advanced progress billing, retention guarantees, and public contracts are only partly supported.",
        "Existing BTP marketing creates acquisition-channel and selection ambiguity.",
    ):
        add_cell_text(right, "• " + item, size=8.4)
    set_table_borders(table, color=WHITE, size="7")

    add_heading(document, "Ideal customer profile", level=2)
    add_body(
        document,
        "Independent tradespeople and small crews, digitally operated, with relatively simple projects, repeat invoicing and expense-management needs, and moderate direct-financing requirements. Larger contractors, long-project businesses, public-works specialists, and firms dependent on overdrafts are a less certain fit.",
        size=9.2,
    )
    add_source_line(document, [
        ("INSEE base", "https://www.insee.fr/fr/statistiques/2011101?geo=FRANCE-1-1"),
        ("FFB", "https://www.ffbatiment.fr/le-batiment-en-chiffres"),
        ("CAPEB", "https://www.capeb.fr/www/capeb/media/national/note-1-trimestre-2026-v3.pdf"),
        ("Shine BTP", "https://www.shine.fr/btp/"),
        ("BTP invoicing limits", "https://help.shine.fr/shine-facture/fr/articles/16596228-artisans-du-btp-facturer-l-avancement-de-vos-chantiers-situations-de-travaux"),
    ])

    add_heading(document, "4. Broader market and regulatory trigger", level=1)
    add_body(
        document,
        "France recorded 1.166 million company creations in 2025, including approximately 759,000 micro-entrepreneurs. Construction created 85,600 businesses but declined 3.8%, while faster-growing sectors included commerce and information and communication. BTP should therefore not consume the entire growth narrative.",
        size=8.8,
    )
    add_body(
        document,
        "From 1 September 2026, all French businesses must receive electronic invoices; from 1 September 2027, SMEs and microbusinesses must issue them. Shine is an accredited platform and includes the service in every plan. Compliance is table stakes; the opportunity is the connected workflow—invoice, collect, reconcile, monitor cash flow, and prepare accounting.",
        size=8.8,
    )
    add_source_line(document, [
        ("INSEE creations", "https://www.insee.fr/fr/statistiques/8721354"),
        ("Government calendar", "https://www.economie.gouv.fr/tout-savoir-sur-la-facturation-electronique-pour-les-entreprises"),
        ("Shine accreditation", "https://help.shine.fr/fr/articles/11459160-shine-plateforme-agreee-de-facturation-electronique"),
    ])


def build_page_four(document):
    add_heading(document, "5. Revised ranked recommendation", level=1, space_before=0)
    initiatives = [
        (
            "1",
            "BTP-fit activation & monetization",
            "Test an integrated workflow around e-invoicing, collection, expense controls, cash-flow signals, and appropriate partner services among suitable tradespeople and small crews.",
            "Randomise against business-as-usual. Scale within BTP only if incremental contribution exceeds treatment and servicing cost with no material guardrail deterioration. Test transferability before expanding to other personas.",
        ),
        (
            "2",
            "Revenue-decline & first-gap protection",
            "Use a material revenue decline as a diagnostic trigger and the first missing-revenue month as the practical intervention window. Monitor recovered accounts separately.",
            "Keep the behavior-based definition: decline is not churn, company closure is not automatically churn, and expected recoverable revenue cannot be inferred without a causal test.",
        ),
    ]
    for number, title, action, rule in initiatives:
        table = document.add_table(rows=1, cols=3)
        table.autofit = False
        for idx, width in enumerate((Cm(1.1), Cm(5.0), Cm(12.2))):
            table.columns[idx].width = width
        cells = table.rows[0].cells
        for cell in cells:
            set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
        set_cell_shading(cells[0], TEAL if number == "1" else NAVY)
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, number, bold=True, color=WHITE, size=16)
        set_cell_shading(cells[1], LIGHT_BLUE)
        add_cell_text(cells[1], title, bold=True, color=NAVY, size=9.2)
        add_cell_text(cells[1], action, size=8.0, color=BLACK, space_after=0)
        add_cell_text(cells[2], "Decision rule", bold=True, color=TEAL, size=8.5)
        add_cell_text(cells[2], rule, size=8.1, color=BLACK, space_after=0)
        set_table_borders(table, color=WHITE, size="7")
        document.add_paragraph().paragraph_format.space_after = Pt(0)

    add_heading(document, "Business initial plan: secondary, not the lead", level=2)
    add_body(
        document,
        "The internal group is small and weak on continuity, the current EUR 60 offer is team-oriented, and competition is strong. Test Business only among companies with observable team and expense-management needs. Obtain effective-dated current-plan history before attributing an effect to the initial plan.",
        size=9.0,
    )

    add_heading(document, "Validate before leadership acts", level=1)
    validations = [
        ("Acquisition & selection", "BTP channel, campaign spend, applicant mix, KYB approval, conversion, CAC, payback, and current penetration."),
        ("Product fit", "Current plan and migration, invoicing, cards, subaccounts, collections, financing requests, rejected payments, and primary-versus-secondary-account use."),
        ("Unit economics", "Gross margin by revenue type, partner economics, support and operational cost, and incremental contribution."),
        ("Governance", "Exact BTP definition and provenance, permitted KYB use, privacy controls, access, retention, and small-cell suppression."),
        ("Causality", "Pre-register the intervention, randomise treatment, use intention-to-treat, and separate association from incremental impact."),
    ]
    for title, detail in validations:
        add_bullet(document, f"{title}: {detail}", size=8.4)

    add_heading(document, "Presentation-ready wording", level=1)
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    add_cell_text(
        cell,
        "BTP is Shine's strongest demonstrated vertical, not yet a proven untapped market. Its large artisan base and workflow fit support a targeted experiment, while sector contraction, financing constraints, advanced invoicing needs, and existing BTP marketing limit broad extrapolation. Test a BTP-specific activation and health journey, using electronic invoicing as the workflow entry point, and scale only when incremental contribution and guardrails outperform business-as-usual.",
        bold=True,
        color=WHITE,
        size=9.4,
        space_after=0,
    )
    set_table_borders(table, color=NAVY, size="1")


def build_references(document):
    add_heading(document, "Linked reference sources", level=1, space_before=0)
    references = [
        ("Shine pricing", "https://www.shine.fr/tarifs/"),
        ("Shine BTP offer", "https://www.shine.fr/btp/"),
        ("Shine: payment institution versus a traditional bank", "https://help.shine.fr/fr/articles/1175387-quelles-sont-les-differences-entre-shine-et-une-banque-dite-traditionnelle"),
        ("Shine overdraft limitation and financing partners", "https://help.shine.fr/fr/articles/1179200-compte-a-decouvert"),
        ("Shine advanced BTP invoicing limitations", "https://help.shine.fr/shine-facture/fr/articles/16596228-artisans-du-btp-facturer-l-avancement-de-vos-chantiers-situations-de-travaux"),
        ("Shine joining Ageras", "https://help.shine.fr/fr/articles/10289941-shine-rejoint-ageras"),
        ("INSEE active legal units by sector", "https://www.insee.fr/fr/statistiques/2011101?geo=FRANCE-1-1"),
        ("INSEE company creations in 2025", "https://www.insee.fr/fr/statistiques/8721354"),
        ("CAPEB Q1 2026 economic outlook", "https://www.capeb.fr/www/capeb/media/national/note-1-trimestre-2026-v3.pdf"),
        ("FFB building-sector figures", "https://www.ffbatiment.fr/le-batiment-en-chiffres"),
        ("French electronic-invoicing reform", "https://www.economie.gouv.fr/tout-savoir-sur-la-facturation-electronique-pour-les-entreprises"),
        ("Official accredited-platform information", "https://www.impots.gouv.fr/facturation-electronique-et-plateformes-agreees"),
        ("Shine accredited-platform confirmation", "https://help.shine.fr/fr/articles/11459160-shine-plateforme-agreee-de-facturation-electronique"),
        ("Qonto pricing", "https://qonto.com/fr/pricing"),
        ("Indy professional account", "https://www.indy.fr/compte-pro/"),
        ("Revolut Business pricing", "https://www.revolut.com/fr-FR/business/business-account-plans/"),
    ]
    for index, (label, url) in enumerate(references, start=1):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.25)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, f"{index}. ", bold=True, color=NAVY, size=8.4)
        add_hyperlink(p, label, url, size=8.4)
        add_run(p, f"  {url}", color=MID_GREY, size=6.8)


def build_document():
    document = Document()
    configure_document(document)
    document.core_properties.title = "Shine external market research and revised strategy"
    document.core_properties.subject = "France market, competition, BTP scalability, and recommendations"
    build_page_one(document)
    document.add_page_break()
    build_page_two(document)
    document.add_page_break()
    build_page_three(document)
    document.add_page_break()
    build_page_four(document)
    document.add_page_break()
    build_references(document)
    document.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    print(build_document())
