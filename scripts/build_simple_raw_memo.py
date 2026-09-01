"""Build the Shine memo using only simple base-table queries."""

from pathlib import Path

import duckdb
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from build_funnel_90d_heatmap import build_chart as build_funnel_heatmap
from build_final_doc import (
    AMBER,
    BLACK,
    LIGHT_BLUE,
    LIGHT_GREY,
    MID_GREY,
    NAVY,
    RED,
    TEAL,
    WHITE,
    add_body,
    add_bullet,
    add_cell_text,
    add_heading,
    add_kpi_strip,
    add_page_title,
    add_run,
    configure_document,
    prevent_row_split,
    set_cell_margins,
    set_cell_shading,
    set_repeat_table_header,
    set_table_borders,
)


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DATABASE = PROJECT_ROOT / "data" / "shine_case.duckdb"
ASSET_DIR = PROJECT_ROOT / "outputs" / "presentation" / "assets"
OUTPUT_FILE = (
    PROJECT_ROOT
    / "outputs"
    / "presentation"
    / "Shine_Case_Simple_RAW_3_Page_Memo.docx"
)


def query_frames():
    connection = duckdb.connect(str(DATABASE), read_only=True)
    try:
        april_persona = connection.execute(
            """
            SELECT
                c.persona,
                COUNT(DISTINCT r.company_profile_id) AS revenue_companies,
                SUM(r.total_revenue) AS total_revenue,
                100.0 * SUM(r.total_revenue)
                    / SUM(SUM(r.total_revenue)) OVER () AS revenue_share_pct,
                SUM(r.total_revenue)
                    / COUNT(DISTINCT r.company_profile_id)
                    AS revenue_per_revenue_company
            FROM companies c
            JOIN revenue_with_total r USING (company_profile_id)
            WHERE r.revenue_month = DATE '2026-04-01'
            GROUP BY c.persona
            ORDER BY total_revenue DESC
            """
        ).fetchdf()
        cumulative_persona = connection.execute(
            """
            SELECT
                c.persona,
                COUNT(DISTINCT r.company_profile_id) AS revenue_companies,
                SUM(r.total_revenue) AS total_revenue,
                100.0 * SUM(r.total_revenue)
                    / SUM(SUM(r.total_revenue)) OVER () AS revenue_share_pct,
                SUM(r.total_revenue)
                    / COUNT(DISTINCT r.company_profile_id)
                    AS revenue_per_revenue_company
            FROM companies c
            JOIN revenue_with_total r USING (company_profile_id)
            WHERE r.revenue_month <= DATE '2026-04-01'
            GROUP BY c.persona
            ORDER BY total_revenue DESC
            """
        ).fetchdf()
        persona_revenue_per_company = connection.execute(
            """
            SELECT
                c.persona,
                COUNT(DISTINCT r.company_profile_id) AS revenue_companies,
                SUM(r.subscription_revenue)
                    / COUNT(DISTINCT r.company_profile_id)
                    AS subscription_revenue_per_company,
                SUM(r.interchange_revenue)
                    / COUNT(DISTINCT r.company_profile_id)
                    AS interchange_revenue_per_company,
                SUM(r.banking_fees)
                    / COUNT(DISTINCT r.company_profile_id)
                    AS banking_fees_per_company,
                SUM(r.deposit_interest_revenue)
                    / COUNT(DISTINCT r.company_profile_id)
                    AS deposit_interest_revenue_per_company,
                SUM(r.total_revenue)
                    / COUNT(DISTINCT r.company_profile_id)
                    AS total_revenue_per_company,
                SUM(r.total_revenue) AS total_revenue
            FROM companies c
            JOIN revenue_with_total r USING (company_profile_id)
            WHERE r.revenue_month <= DATE '2026-04-01'
            GROUP BY c.persona
            ORDER BY total_revenue DESC
            """
        ).fetchdf()
        plan_revenue_per_company = connection.execute(
            """
            SELECT
                c.initial_subscription_group AS initial_plan,
                COUNT(DISTINCT r.company_profile_id) AS revenue_companies,
                SUM(r.subscription_revenue)
                    / COUNT(DISTINCT r.company_profile_id)
                    AS subscription_revenue_per_company,
                SUM(r.interchange_revenue)
                    / COUNT(DISTINCT r.company_profile_id)
                    AS interchange_revenue_per_company,
                SUM(r.banking_fees)
                    / COUNT(DISTINCT r.company_profile_id)
                    AS banking_fees_per_company,
                SUM(r.deposit_interest_revenue)
                    / COUNT(DISTINCT r.company_profile_id)
                    AS deposit_interest_revenue_per_company,
                SUM(r.total_revenue)
                    / COUNT(DISTINCT r.company_profile_id)
                    AS total_revenue_per_company,
                SUM(r.total_revenue) AS total_revenue
            FROM companies c
            JOIN revenue_with_total r USING (company_profile_id)
            WHERE r.revenue_month < DATE '2026-05-01'
              AND c.company_signup_at < DATE '2026-05-01'
            GROUP BY c.initial_subscription_group
            ORDER BY total_revenue_per_company DESC
            """
        ).fetchdf()
        april_plan = connection.execute(
            """
            SELECT
                c.initial_subscription_group AS initial_plan,
                COUNT(DISTINCT r.company_profile_id) AS revenue_companies,
                SUM(r.total_revenue) AS total_revenue,
                100.0 * SUM(r.total_revenue)
                    / SUM(SUM(r.total_revenue)) OVER () AS revenue_share_pct,
                SUM(r.total_revenue)
                    / COUNT(DISTINCT r.company_profile_id)
                    AS revenue_per_revenue_company
            FROM companies c
            JOIN revenue_with_total r USING (company_profile_id)
            WHERE r.revenue_month = DATE '2026-04-01'
            GROUP BY c.initial_subscription_group
            ORDER BY total_revenue DESC
            """
        ).fetchdf()
        cumulative_plan = connection.execute(
            """
            SELECT
                c.initial_subscription_group AS initial_plan,
                COUNT(DISTINCT r.company_profile_id) AS revenue_companies,
                SUM(r.total_revenue) AS total_revenue,
                100.0 * SUM(r.total_revenue)
                    / SUM(SUM(r.total_revenue)) OVER () AS revenue_share_pct,
                SUM(r.total_revenue)
                    / COUNT(DISTINCT r.company_profile_id)
                    AS revenue_per_revenue_company
            FROM companies c
            JOIN revenue_with_total r USING (company_profile_id)
            WHERE r.revenue_month <= DATE '2026-04-01'
            GROUP BY c.initial_subscription_group
            ORDER BY total_revenue DESC
            """
        ).fetchdf()
        april_mix = connection.execute(
            """
            SELECT
                SUM(subscription_revenue) AS subscription_revenue,
                SUM(interchange_revenue) AS interchange_revenue,
                SUM(banking_fees) AS banking_fees,
                SUM(deposit_interest_revenue) AS deposit_interest_revenue,
                SUM(total_revenue) AS total_revenue,
                COUNT(DISTINCT company_profile_id) AS revenue_companies
            FROM revenue_with_total
            WHERE revenue_month = DATE '2026-04-01'
            """
        ).fetchdf().iloc[0]
        confirmed_summary = connection.execute(
            """
            SELECT
                SUM(total_revenue) AS total_revenue,
                COUNT(DISTINCT company_profile_id) AS revenue_companies
            FROM revenue_with_total
            WHERE revenue_month <= DATE '2026-04-01'
            """
        ).fetchdf().iloc[0]
        monthly_revenue = connection.execute(
            """
            SELECT
                revenue_month,
                COUNT(DISTINCT company_profile_id) AS revenue_companies,
                SUM(total_revenue) AS total_revenue
            FROM revenue_with_total
            GROUP BY revenue_month
            ORDER BY revenue_month
            """
        ).fetchdf()
        decline = connection.execute(
            """
            WITH company_revenue AS (
                SELECT
                    company_profile_id,
                    SUM(CASE WHEN revenue_month = DATE '2026-03-01'
                        THEN total_revenue ELSE 0 END) AS march_revenue,
                    SUM(CASE WHEN revenue_month = DATE '2026-04-01'
                        THEN total_revenue ELSE 0 END) AS april_revenue
                FROM revenue_with_total
                WHERE revenue_month IN (
                    DATE '2026-03-01', DATE '2026-04-01'
                )
                GROUP BY company_profile_id
            )
            SELECT
                COUNT(*) FILTER (
                    WHERE march_revenue > 0 AND april_revenue > 0
                ) AS active_in_both_months,
                COUNT(*) FILTER (
                    WHERE march_revenue > 0
                      AND april_revenue > 0
                      AND april_revenue <= 0.70 * march_revenue
                      AND march_revenue - april_revenue >= 10
                ) AS material_decline_companies,
                100.0 * COUNT(*) FILTER (
                    WHERE march_revenue > 0
                      AND april_revenue > 0
                      AND april_revenue <= 0.70 * march_revenue
                      AND march_revenue - april_revenue >= 10
                ) / COUNT(*) FILTER (
                    WHERE march_revenue > 0 AND april_revenue > 0
                ) AS material_decline_share_pct,
                COUNT(*) FILTER (
                    WHERE march_revenue > 0 AND april_revenue = 0
                ) AS first_observed_gap_in_april
            FROM company_revenue
            """
        ).fetchdf().iloc[0]
        health_sql = (
            PROJECT_ROOT / "sql" / "raw_two_table_health_analysis.sql"
        ).read_text(encoding="utf-8")
        health_by_segment = connection.execute(health_sql).fetchdf()
        return (
            april_persona,
            cumulative_persona,
            persona_revenue_per_company,
            plan_revenue_per_company,
            april_plan,
            cumulative_plan,
            april_mix,
            confirmed_summary,
            monthly_revenue,
            decline,
            health_by_segment,
        )
    finally:
        connection.close()


def clean_label(value):
    label = str(value).replace("_", " ").title()
    return label.replace("Btp", "BTP").replace(" It", " IT")


def make_segment_chart(cumulative_persona, cumulative_plan, output_path):
    persona = cumulative_persona.head(8).iloc[::-1].copy()
    plan = cumulative_plan.iloc[::-1].copy()
    fig, axes = plt.subplots(1, 2, figsize=(12.4, 4.7), gridspec_kw={"width_ratios": [1.45, 1]})
    fig.patch.set_facecolor("white")

    bars = axes[0].barh(
        [clean_label(value) for value in persona["persona"]],
        persona["revenue_share_pct"],
        color="#A9C7C7",
        height=0.62,
    )
    for bar, share, average in zip(
        bars,
        persona["revenue_share_pct"],
        persona["revenue_per_revenue_company"],
    ):
        axes[0].text(
            bar.get_width() + 0.35,
            bar.get_y() + bar.get_height() / 2,
            f"{share:.1f}%  |  €{average:.0f}/company",
            va="center",
            fontsize=8.3,
            fontweight="bold",
            color="#17212B",
        )
    axes[0].set_title("Confirmed revenue share by persona", loc="left", fontsize=12, fontweight="bold", color="#102A43")
    axes[0].set_xlabel("Share of cumulative revenue through April", fontsize=9, color="#5E6C76")
    axes[0].set_xlim(0, max(persona["revenue_share_pct"]) * 1.70)

    bars = axes[1].barh(
        [clean_label(value) for value in plan["initial_plan"]],
        plan["revenue_share_pct"],
        color="#A9C7C7",
        height=0.58,
    )
    for bar, share, average in zip(
        bars,
        plan["revenue_share_pct"],
        plan["revenue_per_revenue_company"],
    ):
        axes[1].text(
            bar.get_width() + 0.45,
            bar.get_y() + bar.get_height() / 2,
            f"{share:.1f}%  |  €{average:.0f}/company",
            va="center",
            fontsize=8.3,
            fontweight="bold",
            color="#17212B",
        )
    axes[1].set_title("Confirmed revenue by initial plan", loc="left", fontsize=12, fontweight="bold", color="#102A43")
    axes[1].set_xlabel("Share of cumulative revenue through April", fontsize=9, color="#5E6C76")
    axes[1].set_xlim(0, max(plan["revenue_share_pct"]) * 1.72)

    for axis in axes:
        axis.spines[["top", "right", "left"]].set_visible(False)
        axis.grid(axis="x", color="#E4E9ED", linewidth=0.7)
        axis.set_axisbelow(True)
        axis.tick_params(axis="y", length=0, labelsize=8.5)
        axis.tick_params(axis="x", labelsize=8, colors="#5E6C76")

    plt.tight_layout(rect=(0, 0.02, 1, 1))
    fig.savefig(output_path, dpi=190, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def make_revenue_per_company_heatmap(
    segment_revenue_per_company,
    segment_column,
    chart_title,
    output_path,
    figure_height=8.0,
):
    heatmap = segment_revenue_per_company.sort_values(
        "total_revenue_per_company", ascending=False
    ).copy()
    value_columns = [
        "subscription_revenue_per_company",
        "interchange_revenue_per_company",
        "banking_fees_per_company",
        "deposit_interest_revenue_per_company",
    ]
    column_labels = [
        "Subscription",
        "Interchange",
        "Banking fees",
        "Deposit interest",
    ]
    values = heatmap[value_columns].to_numpy()
    totals = heatmap["total_revenue_per_company"].to_numpy()
    row_labels = [
        f"{clean_label(getattr(row, segment_column))}  "
        f"(n={int(row.revenue_companies):,})"
        for row in heatmap.itertuples()
    ]

    fig, axis = plt.subplots(figsize=(11.3, figure_height))
    fig.patch.set_facecolor("white")
    image = axis.imshow(values, cmap="GnBu", aspect="auto")
    total_x = 4.25
    axis.set_xticks(
        [0, 1, 2, 3, total_x],
        [*column_labels, "Total"],
    )
    axis.set_xlim(-0.5, 4.65)
    axis.get_xticklabels()[-1].set_fontweight("bold")
    axis.get_xticklabels()[-1].set_color("#102A43")
    axis.set_yticks(range(len(row_labels)), row_labels)
    axis.tick_params(
        axis="x",
        labelrotation=0,
        labelsize=10,
        length=0,
        pad=8,
        top=True,
        labeltop=True,
        bottom=False,
        labelbottom=False,
    )
    axis.tick_params(axis="y", labelsize=8.5, length=0, pad=7)
    axis.set_title(
        chart_title,
        loc="left",
        fontsize=14,
        fontweight="bold",
        color="#102A43",
        pad=44,
    )
    axis.set_xlabel(
        "Revenue component — confirmed period through April 2026",
        fontsize=9,
        color="#5E6C76",
        labelpad=10,
    )

    threshold = (values.min() + values.max()) / 2
    for row_index in range(values.shape[0]):
        for column_index in range(values.shape[1]):
            value = values[row_index, column_index]
            axis.text(
                column_index,
                row_index,
                f"€{value:.0f}",
                ha="center",
                va="center",
                fontsize=8.7,
                fontweight="bold",
                color="white" if value > threshold else "#17212B",
            )
        axis.text(
            total_x,
            row_index,
            f"€{totals[row_index]:.0f}",
            ha="center",
            va="center",
            fontsize=9.2,
            fontweight="bold",
            color="#102A43",
        )

    axis.axvline(3.62, color="#A9B4BC", linewidth=1.0)

    colorbar = fig.colorbar(image, ax=axis, fraction=0.035, pad=0.025)
    colorbar.set_label("Cumulative € per revenue company", fontsize=9)
    colorbar.ax.tick_params(labelsize=8)
    colorbar.outline.set_visible(False)
    for spine in axis.spines.values():
        spine.set_visible(False)

    fig.text(
        0.01,
        0.005,
        "n = companies with at least one revenue row through April. Rows are sorted by Total descending; Total is excluded from the heatmap colour scale.",
        fontsize=8,
        color="#5E6C76",
    )
    plt.tight_layout(rect=(0, 0.035, 1, 1))
    fig.savefig(output_path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def make_two_table_health_chart(health_by_segment, output_path):
    states = [
        "Healthy / stable",
        "Declining - low exposure",
        "Material Watch",
        "Active - building history",
        "Recovered / monitor",
        "At-risk",
        "Churned proxy",
        "Never monetized",
    ]
    state_labels = {
        "Healthy / stable": "Healthy / stable",
        "Declining - low exposure": "Declining <€10",
        "Material Watch": "Material Watch",
        "Active - building history": "Building history",
        "Recovered / monitor": "Recovered",
        "At-risk": "At-risk",
        "Churned proxy": "Churned proxy",
        "Never monetized": "Never monetized",
    }
    state_colors = {
        "Healthy / stable": "#16877A",
        "Declining - low exposure": "#F5D58A",
        "Material Watch": "#E9A23B",
        "Active - building history": "#8CB8A8",
        "Recovered / monitor": "#6AAFC1",
        "At-risk": "#D97941",
        "Churned proxy": "#A5453F",
        "Never monetized": "#D8DEE3",
    }
    risk_states = [
        "Declining - low exposure",
        "Material Watch",
        "At-risk",
        "Churned proxy",
    ]

    fig = plt.figure(figsize=(12.4, 8.6), facecolor="white")
    grid = fig.add_gridspec(
        2,
        2,
        height_ratios=[1.25, 2.0],
        width_ratios=[1.35, 1.0],
        hspace=0.42,
        wspace=0.42,
    )
    trend_axis = fig.add_subplot(grid[0, :])
    persona_axis = fig.add_subplot(grid[1, 0])
    plan_axis = fig.add_subplot(grid[1, 1])

    def draw_stacked(axis, frame, index_order, labels, title, show_risk=False):
        pivot = frame.pivot_table(
            index="segment_key",
            columns="account_health_state",
            values="company_share_pct",
            aggfunc="sum",
            fill_value=0,
        ).reindex(index_order).fillna(0)
        counts = frame.groupby("segment_key")["companies"].sum()
        left = [0.0] * len(index_order)
        handles = []
        for state in states:
            values = (
                pivot[state].tolist()
                if state in pivot.columns
                else [0.0] * len(index_order)
            )
            bars = axis.barh(
                range(len(index_order)),
                values,
                left=left,
                color=state_colors[state],
                height=0.62,
                label=state_labels[state],
            )
            handles.append(bars[0])
            if state == "Healthy / stable":
                for row_index, (start, value) in enumerate(zip(left, values)):
                    if value >= 12:
                        axis.text(
                            start + value / 2,
                            row_index,
                            f"{value:.0f}%",
                            ha="center",
                            va="center",
                            color="white",
                            fontsize=7.5,
                            fontweight="bold",
                        )
            left = [start + value for start, value in zip(left, values)]

        axis.set_yticks(range(len(index_order)), labels)
        axis.set_xlim(0, 112 if show_risk else 105)
        axis.set_xticks([0, 25, 50, 75, 100])
        axis.set_xlabel("Share of companies in view", fontsize=8, color="#5E6C76")
        axis.set_title(title, loc="left", fontsize=11, fontweight="bold", color="#102A43")
        axis.grid(axis="x", color="#E4E9ED", linewidth=0.7)
        axis.set_axisbelow(True)
        axis.spines[["top", "right", "left"]].set_visible(False)
        axis.tick_params(axis="y", length=0, labelsize=7.6)
        axis.tick_params(axis="x", labelsize=7.4, colors="#5E6C76")

        if show_risk:
            for row_index, key in enumerate(index_order):
                risk = sum(
                    float(pivot.loc[key, state])
                    if state in pivot.columns
                    else 0.0
                    for state in risk_states
                )
                axis.text(
                    101.0,
                    row_index,
                    f"risk {risk:.0f}%",
                    va="center",
                    fontsize=7.0,
                    fontweight="bold",
                    color="#A5453F",
                )
        return handles

    trend = health_by_segment.loc[
        health_by_segment["segment_level"] == "fixed_cohort_trend"
    ].copy()
    trend["segment_key"] = trend["observation_month"].dt.strftime("%b %y")
    trend_order = (
        trend[["observation_month", "segment_key"]]
        .drop_duplicates()
        .sort_values("observation_month")["segment_key"]
        .tolist()
    )
    trend_counts = trend.groupby("segment_key")["companies"].sum()
    trend_labels = [f"{key}  (n={int(trend_counts[key]):,})" for key in trend_order]
    handles = draw_stacked(
        trend_axis,
        trend,
        trend_order,
        trend_labels,
        "Fixed October activation cohort: all confirmed months",
    )

    april = health_by_segment.loc[
        health_by_segment["observation_month"] == health_by_segment["observation_month"].max()
    ].copy()
    persona = april.loc[april["segment_level"] == "mature_persona"].copy()
    persona["segment_key"] = persona["segment"]
    persona_counts = persona.groupby("segment_key")["companies"].sum()
    largest_personas = persona_counts.nlargest(8).index.tolist()
    persona = persona.loc[persona["segment_key"].isin(largest_personas)]
    persona_risk = (
        persona.loc[persona["account_health_state"].isin(risk_states)]
        .groupby("segment_key")["company_share_pct"]
        .sum()
    )
    persona_order = sorted(
        largest_personas,
        key=lambda value: float(persona_risk.get(value, 0)),
    )
    persona_labels = [
        f"{clean_label(value)}  (n={int(persona_counts[value]):,})"
        for value in persona_order
    ]
    draw_stacked(
        persona_axis,
        persona,
        persona_order,
        persona_labels,
        "April mature accounts: eight largest personas",
        show_risk=True,
    )

    plan = april.loc[april["segment_level"] == "mature_initial_plan"].copy()
    plan["segment_key"] = plan["segment"]
    plan_counts = plan.groupby("segment_key")["companies"].sum()
    plan_risk = (
        plan.loc[plan["account_health_state"].isin(risk_states)]
        .groupby("segment_key")["company_share_pct"]
        .sum()
    )
    plan_order = sorted(
        plan_counts.index.tolist(),
        key=lambda value: float(plan_risk.get(value, 0)),
    )
    plan_labels = [
        f"{clean_label(value)}  (n={int(plan_counts[value]):,})"
        for value in plan_order
    ]
    draw_stacked(
        plan_axis,
        plan,
        plan_order,
        plan_labels,
        "April mature accounts: initial plan",
        show_risk=True,
    )

    fig.suptitle(
        "Revenue-based account health — full confirmed period",
        x=0.01,
        y=0.995,
        ha="left",
        fontsize=14,
        fontweight="bold",
        color="#102A43",
    )
    fig.legend(
        handles,
        [state_labels[state] for state in states],
        ncol=4,
        loc="upper center",
        bbox_to_anchor=(0.5, 0.965),
        frameon=False,
        fontsize=7.7,
    )
    fig.text(
        0.01,
        0.008,
        f"Risk = both declining states + At-risk + Churned proxy. October-April trend fixes the activation cohort; April cuts use {int(persona_counts.sum()):,} mature accounts. Building history is not called Healthy or risky. May excluded.",
        fontsize=7.6,
        color="#5E6C76",
    )
    plt.tight_layout(rect=(0, 0.045, 1, 0.92))
    fig.savefig(output_path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def make_trend_decline_chart(monthly_revenue, decline, output_path):
    months = monthly_revenue.copy()
    months["month_label"] = months["revenue_month"].dt.strftime("%b %y")
    months["revenue_thousands"] = months["total_revenue"] / 1000

    active_both = int(decline["active_in_both_months"])
    material_decline = int(decline["material_decline_companies"])
    april_gap = int(decline["first_observed_gap_in_april"])
    active_no_material_decline = active_both - material_decline

    fig, axes = plt.subplots(
        1,
        2,
        figsize=(12.4, 4.4),
        gridspec_kw={"width_ratios": [1.45, 1]},
    )
    fig.patch.set_facecolor("white")

    colors = ["#008C8C"] * (len(months) - 1) + ["#E9A23B"]
    bars = axes[0].bar(
        months["month_label"],
        months["revenue_thousands"],
        color=colors,
        width=0.68,
    )
    bars[-1].set_hatch("///")
    bars[-1].set_edgecolor("#A56B12")
    for index, (bar, value) in enumerate(zip(bars, months["revenue_thousands"])):
        axes[0].text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 10,
            f"€{value:.0f}k",
            ha="center",
            va="bottom",
            fontsize=8.2,
            fontweight="bold",
            color="#A56B12" if index == len(bars) - 1 else "#17212B",
        )
    axes[0].text(
        len(months) - 1,
        months["revenue_thousands"].iloc[-1] * 0.55,
        "PROVISIONAL",
        ha="center",
        va="center",
        rotation=90,
        fontsize=7.5,
        fontweight="bold",
        color="white",
    )
    axes[0].set_title(
        "Monthly recognized revenue",
        loc="left",
        fontsize=12,
        fontweight="bold",
        color="#102A43",
    )
    axes[0].set_ylabel("Revenue (€ thousands)", fontsize=9, color="#5E6C76")
    axes[0].set_ylim(0, max(months["revenue_thousands"]) * 1.18)
    axes[0].tick_params(axis="x", rotation=35)

    labels = [
        "Active both;\nno material decline",
        "Active both;\n≥30% and €10 decline",
        "March revenue;\nno April row",
    ]
    values = [active_no_material_decline, material_decline, april_gap]
    decline_colors = ["#A9C7C7", "#E9A23B", "#C95C3D"]
    decline_bars = axes[1].barh(labels[::-1], values[::-1], color=decline_colors[::-1], height=0.58)
    for bar, value in zip(decline_bars, values[::-1]):
        axes[1].text(
            bar.get_width() + 180,
            bar.get_y() + bar.get_height() / 2,
            f"{value:,}",
            va="center",
            fontsize=9,
            fontweight="bold",
            color="#17212B",
        )
    axes[1].set_title(
        "Simple March → April screen",
        loc="left",
        fontsize=12,
        fontweight="bold",
        color="#102A43",
    )
    axes[1].set_xlabel("Companies", fontsize=9, color="#5E6C76")
    axes[1].set_xlim(0, max(values) * 1.20)

    for axis in axes:
        axis.spines[["top", "right", "left"]].set_visible(False)
        axis.grid(axis="y" if axis is axes[0] else "x", color="#E4E9ED", linewidth=0.7)
        axis.set_axisbelow(True)
        axis.tick_params(axis="y", length=0, labelsize=8.3)
        axis.tick_params(axis="x", labelsize=8, colors="#5E6C76")

    fig.text(
        0.01,
        0.005,
        "May is retained as a provisional sensitivity. The decline screen is diagnostic only: a decline or missing row is not confirmed churn.",
        fontsize=8,
        color="#5E6C76",
    )
    plt.tight_layout(rect=(0, 0.06, 1, 1))
    fig.savefig(output_path, dpi=190, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def add_code_box(document, lines, caption):
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(18.3)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GREY)
    set_cell_margins(cell, top=75, start=115, bottom=75, end=115)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_run(p, caption, bold=True, color=TEAL, size=7.5)
    for index, line in enumerate(lines):
        p = cell.add_paragraph() if index or p.text else p
        p.paragraph_format.space_after = Pt(0)
        add_run(p, line, color=BLACK, size=6.8)
    set_table_borders(table, color="D8DEE3", size="3")


def build_page_one(
    document,
    chart,
    april_persona,
    cumulative_persona,
    cumulative_plan,
    april_mix,
    confirmed_summary,
):
    btp_april = april_persona.loc[april_persona["persona"] == "BTP"].iloc[0]
    btp_cumulative = cumulative_persona.loc[cumulative_persona["persona"] == "BTP"].iloc[0]
    start = cumulative_plan.loc[cumulative_plan["initial_plan"] == "start"].iloc[0]
    plus = cumulative_plan.loc[cumulative_plan["initial_plan"] == "plus"].iloc[0]
    add_page_title(
        document,
        "Page 1 • Raw-data big picture",
        "BTP leads the full confirmed revenue view; Start leads plan scale",
        "Base tables only • Cumulative through April 2026 • May excluded from headlines",
    )
    add_kpi_strip(
        document,
        [
            (f"€{confirmed_summary['total_revenue'] / 1_000_000:.2f}m", "confirmed revenue through April"),
            (f"{int(confirmed_summary['revenue_companies']):,}", "companies with confirmed revenue"),
            (f"{btp_cumulative['revenue_share_pct']:.1f}%", "confirmed revenue from BTP"),
            (f"{start['revenue_share_pct']:.1f}%", "confirmed revenue from Start"),
        ],
    )
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(str(chart), width=Cm(18.0))

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "How to read: ", bold=True, color=TEAL, size=7.7)
    add_run(
        p,
        "bar length is cumulative revenue share through April; labels add cumulative revenue per revenue company.",
        color=MID_GREY,
        size=7.7,
    )

    add_heading(document, "What I conclude", level=2, space_before=2)
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    for idx, width in enumerate((Cm(6.1), Cm(6.1), Cm(6.1))):
        table.columns[idx].width = width
    cards = [
        ("BTP", f"Largest confirmed persona at {btp_cumulative['revenue_share_pct']:.2f}%; April independently shows {btp_april['revenue_share_pct']:.2f}%, so the result is not driven by one view.", TEAL),
        ("Start vs Plus", f"Start supplies {start['revenue_share_pct']:.2f}% through scale; Plus supplies {plus['revenue_share_pct']:.2f}% at €{plus['revenue_per_revenue_company']:.2f} cumulative revenue per revenue company.", NAVY),
        ("Current economics", "In April, interchange is 42.0% and deposit interest 28.7% of revenue; subscriptions alone are 21.0%.", AMBER),
    ]
    for idx, (title, body, color) in enumerate(cards):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BLUE if idx == 0 else LIGHT_GREY)
        set_cell_margins(cell, top=95, start=110, bottom=95, end=110)
        add_cell_text(cell, title, bold=True, color=color, size=8.7)
        add_cell_text(cell, body, size=7.8, space_after=0)
    set_table_borders(table, color=WHITE, size="7")

    add_code_box(
        document,
        [
            "SELECT c.persona, SUM(r.total_revenue),",
            "       100 * SUM(r.total_revenue) / SUM(SUM(r.total_revenue)) OVER ()",
            "FROM companies c JOIN revenue_with_total r USING (company_profile_id)",
            "WHERE r.revenue_month <= DATE '2026-04-01' GROUP BY c.persona;",
        ],
        "Core query pattern",
    )


def add_initiative(document, number, title, evidence, action, limit, color):
    table = document.add_table(rows=1, cols=4)
    table.autofit = False
    for idx, width in enumerate((Cm(1.1), Cm(4.1), Cm(7.0), Cm(6.1))):
        table.columns[idx].width = width
    cells = table.rows[0].cells
    for cell in cells:
        set_cell_margins(cell, top=95, start=105, bottom=95, end=105)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cells[0], color)
    p = cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, str(number), bold=True, color=WHITE, size=16)
    set_cell_shading(cells[1], LIGHT_BLUE if number == 1 else LIGHT_GREY)
    add_cell_text(cells[1], title, bold=True, color=NAVY, size=9.0)
    add_cell_text(cells[1], action, size=7.8, color=MID_GREY, space_after=0)
    add_cell_text(cells[2], "Raw evidence", bold=True, color=TEAL, size=7.8)
    add_cell_text(cells[2], evidence, size=8.0, space_after=0)
    add_cell_text(cells[3], "Important limit", bold=True, color=RED, size=7.8)
    add_cell_text(cells[3], limit, size=7.8, space_after=0)
    set_table_borders(table, color=WHITE, size="7")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def build_page_two(document, trend_decline_chart, april_persona, cumulative_persona, decline):
    btp_april = april_persona.loc[april_persona["persona"] == "BTP"].iloc[0]
    btp_cumulative = cumulative_persona.loc[cumulative_persona["persona"] == "BTP"].iloc[0]
    add_page_title(
        document,
        "Page 2 • Ranked actions",
        "Use simple evidence to choose where to test—not to claim causality",
        "Two initiatives, clearly weighted",
    )
    add_initiative(
        document,
        1,
        "Post-KYB activation test",
        "In the fair 90-day signup funnel, 66.5% validate, 76.1% of validated companies activate, and 50.6% activate overall (N=22,033).",
        "Randomise validated, unclosed, unactivated companies to improved onboarding versus business-as-usual; measure 90-day revenue per validated company.",
        "Closure reasons, treatment cost, margin, product usage, and the achievable causal uplift are missing. Do not relax KYB controls.",
        TEAL,
    )
    add_initiative(
        document,
        2,
        "Revenue-decline diagnosis",
        f"{int(decline['material_decline_companies']):,} of {int(decline['active_in_both_months']):,} March-April active companies ({decline['material_decline_share_pct']:.2f}%) fell ≥30% and €10; {int(decline['first_observed_gap_in_april']):,} had a first April gap.",
        "Use the decline as a diagnostic trigger and investigate the first missing-revenue month; monitor recoveries separately.",
        "Two months are noisy. This is not churn, forecast loss, or proof that outreach will recover revenue.",
        AMBER,
    )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(str(trend_decline_chart), width=Cm(17.6))

    add_heading(document, "Alternatives considered", level=2, space_before=3)
    alternatives = [
        ("Acquire Business", "High revenue/company, but only 262 April revenue companies and 6.79% of April revenue."),
        ("Immediate plan upsell", "Initial plan is not current plan and correlation is not a causal effect."),
        ("Blanket churn campaign", "A decline, revenue gap, or company closure does not prove customer churn."),
        ("Headline May", "May revenue is retained as sensitivity, but signup coverage is incomplete."),
    ]
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.6)
    table.columns[1].width = Cm(13.7)
    for idx, header in enumerate(("Alternative", "Why it is not the lead")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=65, start=90, bottom=65, end=90)
        add_cell_text(cell, header, bold=True, color=WHITE, size=7.8, space_after=0)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(alternatives):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(values):
            set_cell_shading(cells[idx], WHITE if row_index % 2 == 0 else LIGHT_GREY)
            set_cell_margins(cells[idx], top=52, start=85, bottom=52, end=85)
            add_cell_text(cells[idx], value, bold=(idx == 0), color=NAVY if idx == 0 else BLACK, size=7.4, space_after=0)
    set_table_borders(table, color="D8DEE3", size="3")

def build_page_three(document):
    add_page_title(
        document,
        "Page 3 • Transparency and proof",
        "Start simple; use advanced SQL to test whether the story survives",
        "Raw numbers remain the entry point—the model is the validation layer",
    )
    add_heading(document, "What each layer answers", level=2, space_before=0)
    rows = [
        ("Largest persona", "April and cumulative GROUP BY", "Equal-tenure first 3 complete months", "sql/revenue_concentration_analysis.sql"),
        ("Revenue concentration", "Not used to pick the segment", "Top-10/top-20 and Pareto sensitivities", "sql/revenue_concentration_analysis.sql"),
        ("Account health", "March-to-April decline", "Trailing median + continuity states", "sql/presentation_kpis.sql"),
        ("Inactivity", "Missing revenue row", "Return curve and lifecycle states", "sql/kpi_deep_dive.sql"),
        ("Cohorts", "Not answered", "Equal-age cohort transitions", "sql/cohort_state_analysis.sql"),
        ("Loyalty", "Not answered", "Streak persistence and restart", "sql/streak_analysis.sql"),
    ]
    table = document.add_table(rows=1, cols=4)
    table.autofit = False
    for idx, width in enumerate((Cm(3.3), Cm(4.2), Cm(5.0), Cm(5.8))):
        table.columns[idx].width = width
    for idx, header in enumerate(("Question", "Simple view", "Advanced check", "Where")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=70, start=85, bottom=70, end=85)
        add_cell_text(cell, header, bold=True, color=WHITE, size=7.6, space_after=0)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(values):
            set_cell_shading(cells[idx], WHITE if row_index % 2 == 0 else LIGHT_GREY)
            set_cell_margins(cells[idx], top=55, start=80, bottom=55, end=80)
            add_cell_text(cells[idx], value, bold=(idx == 0), color=NAVY if idx == 0 else BLACK, size=7.1, space_after=0)
    set_table_borders(table, color="D8DEE3", size="3")

    add_heading(document, "The useful triangulation", level=2, space_before=3)
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    for idx in range(3):
        table.columns[idx].width = Cm(6.1)
    checks = [
        ("24.06%", "BTP share in April", "simple current snapshot"),
        ("25.42%", "BTP share through April", "simple cumulative view"),
        ("26.74%", "BTP equal-tenure share", "advanced confirmation"),
    ]
    for idx, (number, label, note) in enumerate(checks):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BLUE if idx == 2 else LIGHT_GREY)
        set_cell_margins(cell, top=105, start=110, bottom=105, end=110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, number, bold=True, color=TEAL if idx == 2 else NAVY, size=16)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, label, bold=True, color=BLACK, size=8.0)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, note, color=MID_GREY, size=7.0)
    set_table_borders(table, color=WHITE, size="7")

    add_heading(document, "Lead experiment and validation", level=2, space_before=3)
    add_bullet(document, "Randomise validated, unclosed, unactivated companies to improved onboarding versus business-as-usual; stratify by persona and initial plan.", size=8.0)
    add_bullet(document, "Use 90-day cumulative revenue per validated company as the primary outcome; activation is the leading indicator.", size=8.0)
    add_bullet(document, "Validate closure meaning, funnel eligibility, treatment cost, margin, product use, channel/CAC, and permitted KYB use.", size=8.0)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    add_run(p, "Simple SQL appendix: ", bold=True, color=TEAL, size=7.7)
    add_run(p, "sql/simple_raw_memo_queries.sql", color=BLACK, size=7.7)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    add_run(p, "Advanced SQL: ", bold=True, color=TEAL, size=7.7)
    add_run(p, "revenue_concentration_analysis.sql • presentation_kpis.sql • cohort_state_analysis.sql • streak_analysis.sql • kpi_deep_dive.sql", color=MID_GREY, size=7.4)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "External context: ", bold=True, color=TEAL, size=7.7)
    add_run(p, "docs/research/SHINE_MARKET_RESEARCH_AND_STRATEGY.md", color=MID_GREY, size=7.4)


def build_page_four(document, heatmap_chart):
    add_page_title(
        document,
        "Page 4 • Revenue-component heatmap",
        "BTP combines broad scale with strong cumulative monetization",
        "Simple raw-data view • Cumulative through April 2026 • May excluded",
    )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(heatmap_chart), width=Cm(16.9))

    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(9.1)
    table.columns[1].width = Cm(9.1)
    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, top=85, start=105, bottom=85, end=105)

    set_cell_shading(left, LIGHT_BLUE)
    add_cell_text(left, "How to read it", bold=True, color=TEAL, size=8.6)
    add_cell_text(
        left,
        "Each coloured cell is the cumulative sum of that revenue component divided by distinct companies with at least one revenue row through April. Darker cells mean more cumulative euros per such company.",
        size=7.6,
    )
    add_cell_text(
        left,
        "The neutral Total column is the sum of the four components and determines the descending row order. It is deliberately excluded from the colour scale so it does not wash out the component differences.",
        size=7.6,
    )
    add_cell_text(
        left,
        "Use Page 1 revenue-share bars for segment scale; use this heatmap to compare monetization intensity and identify which revenue types drive it.",
        size=7.6,
        space_after=0,
    )

    set_cell_shading(right, LIGHT_GREY)
    add_cell_text(right, "Limitations", bold=True, color=RED, size=8.6)
    add_cell_text(
        right,
        "This is cumulative revenue per revenue company, not a monthly average. Older companies have had more months to accumulate revenue, so tenure can influence the ranking.",
        size=7.6,
    )
    add_cell_text(
        right,
        "The inner join excludes companies with no revenue row. Small personas can also be volatile—especially Repair Installation (n=38). Revenue is not profit and does not prove that persona membership caused performance.",
        size=7.6,
        space_after=0,
    )
    set_table_borders(table, color=WHITE, size="7")

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Reproducible SQL: ", bold=True, color=TEAL, size=7.8)
    add_run(p, "Query 9 in sql/simple_raw_memo_queries.sql", color=BLACK, size=7.8)


def build_page_five(document, heatmap_chart):
    add_page_title(
        document,
        "Page 5 • Initial-plan revenue heatmap",
        "Business has the highest cumulative revenue per revenue company—but limited scale",
        "Simple raw-data view • Initial plan, not necessarily current plan • May excluded",
    )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(heatmap_chart), width=Cm(16.9))

    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(9.1)
    table.columns[1].width = Cm(9.1)
    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, top=90, start=105, bottom=90, end=105)

    set_cell_shading(left, LIGHT_BLUE)
    add_cell_text(left, "How to read it", bold=True, color=TEAL, size=8.6)
    add_cell_text(
        left,
        "Each coloured cell is cumulative component revenue divided by distinct revenue companies in that initial plan. The neutral Total column sums the four components and sets the descending row order.",
        size=7.7,
    )
    add_cell_text(
        left,
        "Total is excluded from the colour scale. Company count (n) provides the scale check: a high per-company value in a small group is not automatically the largest opportunity.",
        size=7.7,
        space_after=0,
    )

    set_cell_shading(right, LIGHT_GREY)
    add_cell_text(right, "Limitations", bold=True, color=RED, size=8.6)
    add_cell_text(
        right,
        "The field records the initial subscription group and may not represent the current plan. Differences may reflect company age, selection, size, or behaviour—not a causal effect of the plan.",
        size=7.7,
    )
    add_cell_text(
        right,
        "Values are cumulative, so older companies have more exposure. The inner join also excludes companies with no revenue row. This cannot justify an upsell without current-plan, migration, margin, and experiment data.",
        size=7.7,
        space_after=0,
    )
    set_table_borders(table, color=WHITE, size="7")

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Reproducible SQL: ", bold=True, color=TEAL, size=7.8)
    add_run(p, "Query 10 in sql/simple_raw_memo_queries.sql", color=BLACK, size=7.8)


def build_page_six(document, health_chart, health_by_segment):
    latest_month = health_by_segment["observation_month"].max()
    overall = health_by_segment.loc[
        (health_by_segment["observation_month"] == latest_month)
        & (health_by_segment["segment_level"] == "mature_overall")
    ].copy()
    shares = dict(
        zip(overall["account_health_state"], overall["company_share_pct"])
    )
    comparable_companies = int(overall["companies"].sum())
    inactive_share = shares.get("At-risk", 0) + shares.get("Churned proxy", 0)
    declining_share = (
        shares.get("Declining - low exposure", 0)
        + shares.get("Material Watch", 0)
    )
    add_page_title(
        document,
        "Page 6 • Revenue-based account health",
        "Separate inactivity from deterioration while an account is still producing revenue",
        "Advanced method • All confirmed months, October–April • May excluded",
    )
    add_kpi_strip(
        document,
        [
            (f"{shares.get('Healthy / stable', 0):.1f}%", "Healthy / stable in April"),
            (f"{declining_share:.1f}%", "Declining in April"),
            (f"{inactive_share:.1f}%", "At-risk + churn proxy"),
            (f"{comparable_companies:,}", "comparable mature companies"),
        ],
    )

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    add_run(p, "How to read it. ", bold=True, color=TEAL, size=8.0)
    add_run(
        p,
        "At each month-end, the model first asks whether an activated company produced revenue. Active accounts are then separated by momentum and continuity; inactive accounts are separated by the length of the revenue gap. These are behavioural signals, not contractual customer statuses.",
        color=BLACK,
        size=8.0,
    )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(health_chart), width=Cm(17.5))

    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(9.1)
    table.columns[1].width = Cm(9.1)
    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, top=78, start=100, bottom=78, end=100)

    set_cell_shading(left, LIGHT_BLUE)
    add_cell_text(left, "Definitions", bold=True, color=TEAL, size=8.4)
    add_cell_text(
        left,
        "Healthy/stable: active, sufficiently observed, no prior post-revenue gap, and less than 30% below baseline. Declining—low exposure: down 30%+ but less than €10. Material Watch: down 30%+ and €10+.",
        size=7.25,
    )
    add_cell_text(
        left,
        "Building history: active but not yet assessable for momentum. Recovered: active after a post-revenue gap. At-risk: absent for 1–2 months. Churned proxy: absent for 3+ months. Never monetized: no revenue to date.",
        size=7.25,
        space_after=0,
    )

    set_cell_shading(right, LIGHT_GREY)
    add_cell_text(right, "Interpretation limits", bold=True, color=RED, size=8.4)
    add_cell_text(
        right,
        "This is revenue health—not product use, satisfaction, profitability, account status, or confirmed cancellation. A missing row is assumed to mean no observed revenue and must be validated.",
        size=7.25,
    )
    add_cell_text(
        right,
        "All confirmed months are shown. Early snapshots naturally contain more Building-history accounts; momentum states only appear after enough complete history exists. The trend fixes the October cohort, while April segment cuts use mature companies. Initial plan may not be current.",
        size=7.25,
        space_after=0,
    )
    set_table_borders(table, color=WHITE, size="7")

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Self-contained SQL: ", bold=True, color=TEAL, size=7.6)
    add_run(
        p,
        "sql/raw_two_table_health_analysis.sql — reads only companies and revenue_with_total",
        color=BLACK,
        size=7.6,
    )


def build_page_seven(document, funnel_heatmap_chart):
    add_page_title(
        document,
        "Page 7 • Fair 90-day funnel",
        "The largest quantifiable growth lever is post-KYB activation across the business",
        "Simple method • Companies table only • May excluded",
    )

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Method: ", bold=True, color=TEAL, size=8.1)
    add_run(
        p,
        "include signups from 1 October 2025 through 30 January 2026, so every company has a complete 90-day observation window through April. Colours are relative within each segment column—not targets—and N<100 should be treated cautiously.",
        color=BLACK,
        size=8.1,
    )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(funnel_heatmap_chart), width=Cm(15.8))

    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(9.1)
    table.columns[1].width = Cm(9.1)
    evidence, action = table.rows[0].cells
    for cell in (evidence, action):
        set_cell_margins(cell, top=80, start=105, bottom=80, end=105)

    set_cell_shading(evidence, LIGHT_BLUE)
    add_cell_text(evidence, "What the data shows", bold=True, color=TEAL, size=8.5)
    add_cell_text(
        evidence,
        "Overall, 66.5% validate, 76.1% of validated companies activate within the same 90-day signup window, and 50.6% activate overall (N=22,033). Initial-plan order is stable: Start 54.7%, Plus 49.6%, Free 48.2%, and Business 32.7%.",
        size=7.45,
        space_after=0,
    )

    set_cell_shading(action, LIGHT_GREY)
    add_cell_text(action, "What Shine should do", bold=True, color=NAVY, size=8.5)
    add_cell_text(
        action,
        "A/B test improved onboarding for validated, unclosed, unactivated companies versus business-as-usual. Use 90-day cumulative revenue per validated company as the primary outcome. A hypothetical +5 pp persistent lift equals about 172 extra activations and €6.8k monthly revenue per matured cohort; this is a scenario, not a forecast.",
        size=7.45,
        space_after=0,
    )
    set_table_borders(table, color=WHITE, size="7")

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Reproducible SQL: ", bold=True, color=TEAL, size=7.7)
    add_run(
        p,
        "sql/funnel_90d_heatmap_queries.sql • coloured copy: outputs/presentation/Colored_90_Day_Funnel_Queries.html",
        color=BLACK,
        size=7.7,
    )


def build_document():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    (
        april_persona,
        cumulative_persona,
        persona_revenue_per_company,
        plan_revenue_per_company,
        _april_plan,
        cumulative_plan,
        april_mix,
        confirmed_summary,
        monthly_revenue,
        decline,
        health_by_segment,
    ) = query_frames()
    chart = ASSET_DIR / "simple_raw_segment_view.png"
    trend_decline_chart = ASSET_DIR / "simple_raw_trend_and_decline.png"
    heatmap_chart = ASSET_DIR / "simple_raw_revenue_per_company_heatmap.png"
    plan_heatmap_chart = ASSET_DIR / "simple_raw_plan_revenue_per_company_heatmap.png"
    health_chart = ASSET_DIR / "raw_two_table_health_full_period.png"
    funnel_heatmap_chart = ASSET_DIR / "simple_raw_funnel_90d_heatmap.jpg"
    make_segment_chart(cumulative_persona, cumulative_plan, chart)
    make_trend_decline_chart(monthly_revenue, decline, trend_decline_chart)
    make_revenue_per_company_heatmap(
        persona_revenue_per_company,
        "persona",
        "Cumulative revenue per revenue company by persona",
        heatmap_chart,
    )
    make_revenue_per_company_heatmap(
        plan_revenue_per_company,
        "initial_plan",
        "Cumulative revenue per revenue company by initial plan",
        plan_heatmap_chart,
        figure_height=4.7,
    )
    make_two_table_health_chart(health_by_segment, health_chart)
    build_funnel_heatmap()

    document = Document()
    configure_document(document)
    document.core_properties.title = "Shine — Simple raw-data revenue memo"
    document.core_properties.subject = "Simple memo using base tables and straightforward SQL"
    build_page_one(
        document,
        chart,
        april_persona,
        cumulative_persona,
        cumulative_plan,
        april_mix,
        confirmed_summary,
    )
    document.add_page_break()
    build_page_two(
        document,
        trend_decline_chart,
        april_persona,
        cumulative_persona,
        decline,
    )
    document.add_page_break()
    build_page_three(document)
    document.add_page_break()
    build_page_four(document, heatmap_chart)
    document.add_page_break()
    build_page_five(document, plan_heatmap_chart)
    document.add_page_break()
    build_page_six(document, health_chart, health_by_segment)
    document.add_page_break()
    build_page_seven(document, funnel_heatmap_chart)
    document.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    print(build_document())
