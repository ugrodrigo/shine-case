"""Build the final three-page Shine case memo as a Word document."""

from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = PROJECT_ROOT / "outputs" / "presentation" / "assets"
OUTPUT_FILE = PROJECT_ROOT / "outputs" / "presentation" / "Shine_Case_3_Page_Memo.docx"

NAVY = "102A43"
TEAL = "008C8C"
GREEN = "2A9D8F"
AMBER = "E9A23B"
RED = "C95C3D"
LIGHT_BLUE = "EAF4F4"
LIGHT_GREY = "F3F5F7"
MID_GREY = "5E6C76"
WHITE = "FFFFFF"
BLACK = "17212B"


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D8DEE3", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_keep_with_next(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if value and keep is None:
        p_pr.append(OxmlElement("w:keepNext"))


def add_run(paragraph, text, *, bold=False, color=BLACK, size=9.2, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_hyperlink(paragraph, text, url, *, color=TEAL, size=7.2):
    """Add a clickable external hyperlink to a paragraph."""
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    properties.append(run_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(int(size * 2)))
    properties.append(font_size)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_body(document, text, *, bold_prefix=None, size=9.2, space_after=3, color=BLACK):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.04
    if bold_prefix and text.startswith(bold_prefix):
        add_run(paragraph, bold_prefix, bold=True, size=size, color=color)
        add_run(paragraph, text[len(bold_prefix):], size=size, color=color)
    else:
        add_run(paragraph, text, size=size, color=color)
    return paragraph


def add_bullet(document, text, *, level=0, size=8.9, color=BLACK, space_after=1.5):
    paragraph = document.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.left_indent = Cm(0.45 + 0.35 * level)
    paragraph.paragraph_format.first_line_indent = Cm(-0.22)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.0
    add_run(paragraph, text, size=size, color=color)
    return paragraph


def add_heading(document, text, level=1, *, color=NAVY, space_before=4, space_after=3):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.keep_with_next = True
    size = 13.2 if level == 1 else 10.6
    add_run(paragraph, text, bold=True, color=color, size=size)
    return paragraph


def add_page_title(document, kicker: str, title: str, subtitle: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    add_run(p, kicker.upper(), bold=True, color=TEAL, size=8.0)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.keep_with_next = True
    add_run(p, title, bold=True, color=NAVY, size=20)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, subtitle, color=MID_GREY, size=8.8)


def add_callout(document, text: str, *, fill=NAVY, color=WHITE, size=11.4) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(18.6)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=145, start=180, bottom=145, end=180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.03
    add_run(p, text, bold=True, color=color, size=size)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_kpi_strip(document, items) -> None:
    table = document.add_table(rows=1, cols=len(items))
    table.autofit = False
    for idx, (value, label) in enumerate(items):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_GREY if idx % 2 == 0 else LIGHT_BLUE)
        set_cell_margins(cell, top=90, start=65, bottom=90, end=65)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_run(p, value, bold=True, color=TEAL, size=15)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_run(p, label, color=MID_GREY, size=7.5)
    set_table_borders(table, color=WHITE, size="8")


def add_small_label(cell, text, color=TEAL) -> None:
    p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    add_run(p, text.upper(), bold=True, color=color, size=7.5)


def add_cell_text(cell, text, *, bold=False, size=8.5, color=BLACK, space_after=2):
    p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, text, bold=bold, size=size, color=color)
    return p


def add_page_break(document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break()
    p.runs[0]._r.get_or_add_br().set(qn("w:type"), "page")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(paragraph, "SHINE CASE  •  ", bold=True, color=MID_GREY, size=7.2)
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char_1, instr_text, fld_char_2))


def make_concentration_chart(path: Path) -> None:
    from matplotlib.patches import Patch

    fig, ax = plt.subplots(figsize=(7.3, 1.75), dpi=190)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    labels = ["Share of companies", "Share of revenue (months 1-3)"]
    top = [20, 70.24]
    remainder = [80, 29.76]
    y = [1, 0]
    ax.barh(y, top, color="#008C8C", height=0.48)
    ax.barh(y, remainder, left=top, color="#DCE6EA", height=0.48)
    direct_labels = [
        ("20%", "80%"),
        ("70%", "30%"),
    ]
    for yi, value, (high_label, other_label) in zip(y, top, direct_labels):
        ax.text(value / 2, yi, high_label, ha="center", va="center",
                color="white", fontsize=8.4, fontweight="bold")
        ax.text(value + (100 - value) / 2, yi, other_label,
                ha="center", va="center", color="#425466", fontsize=8.1)
    ax.set_yticks(y, labels, fontsize=8.5, color="#17212B")
    ax.set_xlim(0, 100)
    ax.set_xticks([])
    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.tick_params(axis="y", length=0, pad=7)
    ax.set_title(
        "The highest-revenue 20% of companies generate 70% of comparable revenue",
        loc="left",
        fontsize=9.3,
        fontweight="bold",
        color="#102A43",
        pad=8,
    )
    legend = [
        Patch(facecolor="#008C8C", label="Highest-revenue 20% of companies"),
        Patch(facecolor="#DCE6EA", label="Other 80% of companies"),
    ]
    ax.legend(
        handles=legend,
        ncol=2,
        frameon=False,
        loc="lower center",
        bbox_to_anchor=(0.5, -0.30),
        fontsize=7.3,
        handlelength=1.1,
        columnspacing=1.5,
    )
    plt.tight_layout(rect=[0, 0.13, 1, 1], pad=0.4)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def make_segment_ranking_chart(path: Path) -> None:
    from matplotlib.patches import Patch

    ranking = pd.read_csv(
        PROJECT_ROOT / "outputs" / "eda" / "eda_segment_opportunity_ranking.csv"
    )
    personas = (
        ranking[ranking["segment_level"].eq("persona")]
        .nsmallest(8, "revenue_opportunity_rank")
        .sort_values("revenue_opportunity_rank")
    )
    plans = (
        ranking[ranking["segment_level"].eq("initial_plan")]
        .sort_values("revenue_opportunity_rank")
    )

    def bar_color(index: float) -> str:
        if index >= 1.20:
            return "#008C8C"
        if index >= 0.80:
            return "#6AAED6"
        return "#B8C3CA"

    def draw_panel(ax, frame, title, x_limit):
        labels = []
        for value in frame["segment_name"]:
            label = str(value).replace("_", " ").title()
            label = label.replace("Btp", "BTP").replace(" It", " IT")
            labels.append(label)
        values = frame["share_of_all_comparable_revenue_pct"].astype(float).tolist()
        indexes = frame["average_revenue_per_company_index"].astype(float).tolist()
        averages = frame["average_first3_full_month_revenue"].astype(float).tolist()
        populations = frame["eligible_companies"].astype(int).tolist()
        colors = [bar_color(value) for value in indexes]
        y = list(range(len(frame)))
        bars = ax.barh(y, values, color=colors, height=0.62)
        ax.set_yticks(y, labels, fontsize=6.8, color="#17212B")
        ax.invert_yaxis()
        ax.set_xlim(0, x_limit)
        ax.set_xlabel("Share of all comparable revenue", fontsize=6.7, color="#425466")
        ax.tick_params(axis="x", labelsize=6.2, colors="#6B7780", length=0)
        ax.tick_params(axis="y", length=0, pad=4)
        ax.grid(axis="x", color="#E8EDF0", linewidth=0.55)
        ax.set_axisbelow(True)
        ax.set_title(title, loc="left", fontsize=8.8, fontweight="bold", color="#102A43", pad=4)
        for bar, share, population, average, index in zip(
            bars, values, populations, averages, indexes
        ):
            if share < 5:
                ax.text(
                    share + 0.5,
                    bar.get_y() + bar.get_height() / 2,
                    f"{share:.1f}%  |  N={population:,}  |  €{average:.0f}/co  |  {index:.2f}×",
                    ha="left",
                    va="center",
                    fontsize=6.2,
                    color="#425466",
                )
            elif share < 8:
                ax.text(
                    share / 2,
                    bar.get_y() + bar.get_height() / 2,
                    f"{share:.1f}%",
                    ha="center",
                    va="center",
                    fontsize=6.5,
                    fontweight="bold",
                    color="white",
                )
                ax.text(
                    share + 0.7,
                    bar.get_y() + bar.get_height() / 2,
                    f"N={population:,}  |  €{average:.0f}/co  |  {index:.2f}×",
                    ha="left",
                    va="center",
                    fontsize=6.2,
                    color="#425466",
                )
            else:
                ax.text(
                    share - 0.8,
                    bar.get_y() + bar.get_height() / 2,
                    f"{share:.1f}%",
                    ha="right",
                    va="center",
                    fontsize=6.5,
                    fontweight="bold",
                    color="white",
                )
                ax.text(
                    share + 0.7,
                    bar.get_y() + bar.get_height() / 2,
                    f"N={population:,}  |  €{average:.0f}/co  |  {index:.2f}×",
                    ha="left",
                    va="center",
                    fontsize=6.2,
                    color="#425466",
                )
        for spine in ax.spines.values():
            spine.set_visible(False)

    fig, axes = plt.subplots(
        1,
        2,
        figsize=(7.3, 2.25),
        dpi=190,
        gridspec_kw={"width_ratios": [1.35, 1.0]},
    )
    fig.patch.set_facecolor("white")
    for ax in axes:
        ax.set_facecolor("white")
    draw_panel(axes[0], personas, "Persona ranking", 56)
    draw_panel(axes[1], plans, "Initial-plan ranking", 76)
    fig.suptitle(
        "Full-population revenue scale and efficiency point to different choices",
        x=0.01,
        y=0.98,
        ha="left",
        fontsize=9.5,
        fontweight="bold",
        color="#102A43",
    )
    legend = [
        Patch(facecolor="#008C8C", label="Revenue/company ≥1.2× overall"),
        Patch(facecolor="#6AAED6", label="Revenue/company 0.8–1.2×"),
        Patch(facecolor="#B8C3CA", label="Revenue/company <0.8×"),
    ]
    fig.legend(
        handles=legend,
        ncol=3,
        frameon=False,
        loc="lower center",
        bbox_to_anchor=(0.5, 0.005),
        fontsize=6.5,
        handlelength=1.0,
        columnspacing=1.2,
    )
    plt.tight_layout(rect=[0, 0.10, 1, 0.90], pad=0.45, w_pad=1.0)
    fig.savefig(path, facecolor="white")
    plt.close(fig)


def make_health_chart(path: Path) -> None:
    from matplotlib.patches import Patch

    health = pd.read_csv(
        PROJECT_ROOT / "outputs" / "eda" / "eda_comparable_account_health_by_segment.csv"
    )
    persona_names = [
        "BTP",
        "Consultant",
        "Retail",
        "Others",
        "Automobile_Trade_Repair",
        "Developer_IT",
        "Bikers_Drivers",
        "Wholesale",
    ]
    states = [
        "Healthy / stable",
        "Declining - low exposure",
        "Material Watch",
        "Recovered / monitor",
        "At-risk",
        "Churned proxy",
        "Never monetized",
    ]
    state_labels = [
        "Healthy / stable",
        "Declining <€10",
        "Material Watch",
        "Recovered",
        "At-risk",
        "Churned proxy",
        "Never monetized",
    ]
    colors = [
        "#2A9D8F",
        "#F5D58A",
        "#E9A23B",
        "#6AAED6",
        "#E07A3F",
        "#C95C3D",
        "#AAB7C0",
    ]

    def prepare(frame, segment_column, names):
        subset = frame[frame[segment_column].isin(names)].copy()
        pivot = subset.pivot_table(
            index=segment_column,
            columns="account_health_state",
            values="comparable_segment_company_share_pct",
            aggfunc="sum",
            fill_value=0,
        )
        populations = subset.groupby(segment_column)["comparable_segment_companies"].max()
        for state in states:
            if state not in pivot.columns:
                pivot[state] = 0.0
        pivot = pivot[states]
        signal = (
            pivot["Declining - low exposure"]
            + pivot["Material Watch"]
            + pivot["At-risk"]
            + pivot["Churned proxy"]
        )
        signal_counts = (
            subset[
                subset["account_health_state"].isin(
                    [
                        "Declining - low exposure",
                        "Material Watch",
                        "At-risk",
                        "Churned proxy",
                    ]
                )
            ]
            .groupby(segment_column)["companies"]
            .sum()
            .reindex(pivot.index, fill_value=0)
        )
        order = signal.sort_values(ascending=False).index
        return (
            pivot.loc[order],
            populations.loc[order],
            signal.loc[order],
            signal_counts.loc[order],
        )

    def draw_stack(ax, pivot, populations, signal, signal_counts, title):
        y = list(range(len(pivot)))
        left = [0.0] * len(pivot)
        for state, label, color in zip(states, state_labels, colors):
            values = pivot[state].astype(float).tolist()
            bars = ax.barh(y, values, left=left, color=color, height=0.72, label=label)
            for bar, value, start in zip(bars, values, left):
                if value >= 7.0:
                    ax.text(
                        start + value / 2,
                        bar.get_y() + bar.get_height() / 2,
                        f"{value:.0f}%",
                        ha="center",
                        va="center",
                        fontsize=7.0,
                        fontweight="bold",
                        color=(
                            "#543A00"
                            if state in ("Declining - low exposure", "Material Watch")
                            else "white"
                        ),
                    )
            left = [a + b for a, b in zip(left, values)]
        labels = []
        for value in pivot.index:
            label = str(value).replace("_", " ").title()
            label = label.replace("Btp", "BTP").replace(" It", " IT")
            labels.append(f"{label}  (N={int(populations[value]):,})")
        ax.set_yticks(y, labels, fontsize=7.4, color="#17212B")
        ax.invert_yaxis()
        ax.set_xlim(0, 121)
        ax.set_xticks([0, 25, 50, 75, 100], ["0%", "25%", "50%", "75%", "100%"])
        ax.tick_params(axis="x", labelsize=6.8, colors="#6B7780", length=0)
        ax.tick_params(axis="y", length=0, pad=4)
        ax.grid(axis="x", color="#E8EDF0", linewidth=0.5)
        ax.set_axisbelow(True)
        ax.set_title(title, loc="left", fontsize=9.5, fontweight="bold", color="#102A43", pad=3)
        for yi, name in enumerate(pivot.index):
            annotation = f"Declining + inactive {signal[name]:.1f}%"
            ax.text(
                101.2,
                yi,
                annotation,
                va="center",
                ha="left",
                fontsize=6.8,
                color="#425466",
            )
        for spine in ax.spines.values():
            spine.set_visible(False)

    persona_frame = health[health["segment_level"].eq("persona")]
    plan_frame = health[health["segment_level"].eq("initial_plan")]
    overall_frame = health[health["segment_level"].eq("overall")].copy()
    overall_frame["overall_label"] = "All mature companies"
    overall_pivot, overall_n, overall_signal, overall_signal_n = prepare(
        overall_frame,
        "overall_label",
        ["All mature companies"],
    )
    persona_pivot, persona_n, persona_signal, persona_signal_n = prepare(persona_frame, "persona", persona_names)
    plan_pivot, plan_n, plan_signal, plan_signal_n = prepare(
        plan_frame,
        "initial_subscription_group",
        ["start", "plus", "free", "business"],
    )

    fig, axes = plt.subplots(
        3,
        1,
        figsize=(7.6, 4.6),
        dpi=190,
        gridspec_kw={"height_ratios": [0.45, 2.0, 1.1]},
    )
    fig.patch.set_facecolor("white")
    draw_stack(axes[0], overall_pivot, overall_n, overall_signal, overall_signal_n, "Overall")
    draw_stack(axes[1], persona_pivot, persona_n, persona_signal, persona_signal_n, "Personas")
    draw_stack(axes[2], plan_pivot, plan_n, plan_signal, plan_signal_n, "Initial plans")
    fig.suptitle(
        "Mature-company health exposes different weak and strong candidates",
        x=0.01,
        y=0.985,
        ha="left",
        fontsize=10.5,
        fontweight="bold",
        color="#102A43",
    )
    overall_name = overall_signal.index[0]
    fig.text(
        0.01,
        0.943,
        f"Overall Declining + inactive: {overall_signal[overall_name]:.1f}% "
        f"(N={int(overall_signal_n[overall_name]):,})",
        ha="left",
        fontsize=7.4,
        color="#425466",
    )
    legend = [Patch(facecolor=color, label=label) for label, color in zip(state_labels, colors)]
    fig.legend(
        handles=legend,
        ncol=4,
        frameon=False,
        loc="lower center",
        bbox_to_anchor=(0.5, 0.002),
        fontsize=7.2,
        handlelength=1.0,
        columnspacing=0.9,
    )
    plt.tight_layout(rect=[0, 0.13, 1, 0.90], pad=0.5, h_pad=0.7)
    fig.savefig(path, facecolor="white")
    plt.close(fig)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.4)
    section.footer_distance = Cm(0.45)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.02

    for style_name in ("List Bullet", "List Bullet 2"):
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(8.9)

    footer = section.footer
    add_page_number(footer.paragraphs[0])

    document.core_properties.title = "Shine — Revenue Risk & Opportunity"
    document.core_properties.subject = "Senior Data Analyst Case Study"
    document.core_properties.author = "Candidate analysis"


def build_page_one(document: Document, ranking_chart: Path) -> None:
    add_page_title(
        document,
        "Page 1 • Position and magnitude",
        "Revenue deterioration is the clearest near-term risk",
        "Confirmed analysis through April 2026 • May excluded because signup coverage is truncated",
    )
    add_callout(
        document,
        "Among 6,337 mature companies, 1,909 (30.1%) are in Declining + inactive. The immediate "
        "test group is narrower: 953 (15.0%) are either Material Watch or At-risk, where Shine "
        "can intervene before a longer period without revenue.",
    )
    add_kpi_strip(
        document,
        [
            ("6,337", "mature companies assessed"),
            ("30.1%", "Declining + inactive"),
            ("10.6%", "Material Watch • N=673"),
            ("4.4%", "At-risk • N=280"),
        ],
    )
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run()
    run.add_picture(str(ranking_chart), width=Cm(17.7))

    method = document.add_table(rows=1, cols=1)
    method.autofit = False
    method.columns[0].width = Cm(18.3)
    cell = method.cell(0, 0)
    set_cell_margins(cell, top=55, start=100, bottom=55, end=100)
    set_cell_shading(cell, LIGHT_GREY)
    add_cell_text(
        cell,
        "Why this window: revenue is supplied only by calendar month. For a company activated "
        "late in month 0, the table cannot separate pre- and post-activation days. We therefore "
        "exclude month 0 and compare months 1-3, the first three complete calendar months. An "
        "exact 90-day window would require daily revenue.",
        size=7.1,
        color=MID_GREY,
    )
    set_table_borders(method, color="DDE3E8", size="4")

    add_heading(document, "What the data shows", level=2, space_before=2)
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(9.15)
    table.columns[1].width = Cm(9.15)
    opportunity, risk = table.rows[0].cells
    for cell in (opportunity, risk):
        set_cell_margins(cell, top=105, start=130, bottom=105, end=130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(opportunity, LIGHT_BLUE)
    set_cell_shading(risk, LIGHT_GREY)
    add_small_label(opportunity, "Lead opportunity", TEAL)
    add_cell_text(opportunity, "Intervene while an account is still revenue-active but materially declining, or immediately after its first revenue-free month.", bold=True)
    add_cell_text(opportunity, "Prioritise the largest exposures first; use low-cost automation for the long tail.", size=8.2)
    add_small_label(risk, "Risk", RED)
    add_cell_text(risk, "Declining + inactive is a screening group—not one account state and not confirmed churn.", bold=True)
    add_cell_text(risk, "It also contains Declining <€10 and Churned proxy, which should receive lower-cost or lower-priority treatment.", size=8.2)
    set_table_borders(table, color=WHITE, size="8")

    add_heading(document, "How to read the ranking", level=2, space_before=3)
    definitions = document.add_table(rows=2, cols=2)
    definitions.autofit = False
    definitions.columns[0].width = Cm(9.15)
    definitions.columns[1].width = Cm(9.15)
    text = [
        ("Bar length", "share of revenue across the complete 9,088-company base"),
        ("Bar color", "average revenue/company versus overall; teal means at least 1.2×"),
        ("N", "eligible companies with three complete post-activation months"),
        ("No composite score", "leadership can see the scale-efficiency trade-off directly"),
    ]
    for idx, (name, definition) in enumerate(text):
        cell = definitions.cell(idx // 2, idx % 2)
        set_cell_margins(cell, top=55, start=95, bottom=55, end=95)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, f"{name}: ", bold=True, color=NAVY, size=7.8)
        add_run(p, definition, color=MID_GREY, size=7.8)
    set_table_borders(definitions, color="DDE3E8", size="4")

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Decision: ", bold=True, color=TEAL, size=9.0)
    add_run(
        p,
        "lead with a controlled experiment for Material Watch and At-risk companies. Prioritise by "
        "gross exposure and expected recovery—not persona alone—and keep BTP discovery second.",
        size=9.0,
    )


def add_initiative_card(document, rank, title, why, magnitude, limit, fill):
    table = document.add_table(rows=1, cols=4)
    table.autofit = False
    widths = [Cm(1.1), Cm(4.2), Cm(7.4), Cm(5.6)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    cells = table.rows[0].cells
    for cell in cells:
        set_cell_margins(cell, top=72, start=95, bottom=72, end=95)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cells[0], fill)
    p = cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, str(rank), bold=True, color=WHITE, size=16)
    set_cell_shading(cells[1], LIGHT_BLUE if rank == 1 else LIGHT_GREY)
    add_cell_text(cells[1], title, bold=True, size=9.3, color=NAVY)
    add_cell_text(cells[1], why, size=7.8, color=MID_GREY)
    add_small_label(cells[2], "Observable magnitude", TEAL)
    add_cell_text(cells[2], magnitude, size=8.2)
    add_small_label(cells[3], "Limit", RED)
    add_cell_text(cells[3], limit, size=8.0)
    set_table_borders(table, color=WHITE, size="7")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def build_page_two(document: Document, health_chart: Path) -> None:
    add_page_title(
        document,
        "Page 2 • Ranked decisions",
        "Two initiatives, clearly weighted",
        "Ranked by observable scale, value concentration, health, and actionability",
    )
    add_initiative_card(
        document,
        1,
        "Revenue-deterioration & first-gap protection",
        "Diagnose material declines while revenue is still active and support accounts after their first missing month; use state-specific treatments.",
        "Material Watch is 10.6% (N=673) and At-risk is 4.4% (N=280): together 15.0% of mature companies (N=953). Their gross exposure proxy is €51.1k; illustrative 10–20% recovery is €5.1k–€10.2k/month before cost—not a forecast.",
        "Exposure is not recoverable revenue. Product usage, cause, recovery probability, contribution margin, and treatment cost are missing.",
        TEAL,
    )
    add_initiative_card(
        document,
        2,
        "BTP product-fit discovery experiment",
        "Validate which BTP needs and product behaviours drive value before scaling acquisition, workflow activation, or upsell.",
        "BTP contributes 26.74% of comparable revenue (N=1,642; €178.85/company; 1.48× overall), combining attractive scale and efficiency.",
        "No product-use, current-plan, CAC, margin, or acquisition-channel data identifies a causal lever; sector and product constraints cap extrapolation.",
        AMBER,
    )

    external = document.add_paragraph()
    external.paragraph_format.space_before = Pt(1)
    external.paragraph_format.space_after = Pt(1)
    add_run(external, "External pressure test: ", bold=True, color=TEAL, size=7.1)
    add_run(external, "large artisan base and a dedicated Shine offer support fit; BTP contraction, no direct overdraft, and advanced invoicing limits narrow the ideal customer. ", size=7.1, color=MID_GREY)
    add_hyperlink(external, "INSEE", "https://www.insee.fr/fr/statistiques/2011101?geo=FRANCE-1-1", size=7.1)
    add_run(external, " • ", size=7.1, color=MID_GREY)
    add_hyperlink(external, "Shine BTP", "https://www.shine.fr/btp/", size=7.1)
    add_run(external, " • ", size=7.1, color=MID_GREY)
    add_hyperlink(external, "CAPEB", "https://www.capeb.fr/www/capeb/media/national/note-1-trimestre-2026-v3.pdf", size=7.1)
    add_run(external, " • ", size=7.1, color=MID_GREY)
    add_hyperlink(external, "Product limits", "https://help.shine.fr/shine-facture/fr/articles/16596228-artisans-du-btp-facturer-l-avancement-de-vos-chantiers-situations-de-travaux", size=7.1)

    add_heading(document, "Health by segment: activated by Dec 2025 (N=6,337)", level=2, space_before=2)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(str(health_chart), width=Cm(16.5))

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    add_run(p, "Legend definitions", bold=True, color=NAVY, size=7.6)

    legend_definitions = [
        (
            "Healthy / stable",
            "revenue-active in April, uninterrupted since activation, and without a decline of 30% or more.",
            GREEN,
        ),
        (
            "Declining <€10",
            "revenue-active and at least 30% below its prior three-month median, but by less than €10.",
            "C58A1D",
        ),
        (
            "Material Watch",
            "revenue-active and at least 30% below its prior three-month median, with a decline of at least €10.",
            AMBER,
        ),
        (
            "Recovered",
            "revenue-active again after an observed no-revenue month since activation.",
            "6AAED6",
        ),
        (
            "At-risk",
            "previously monetized, with no observed revenue for one or two months.",
            "E07A3F",
        ),
        (
            "Churned proxy",
            "previously monetized, with no observed revenue for at least three months.",
            RED,
        ),
        (
            "Never monetized",
            "no revenue has been observed for the company in the available extract.",
            "7A8B96",
        ),
    ]
    definitions = document.add_table(rows=4, cols=2)
    definitions.autofit = False
    definitions.columns[0].width = Cm(9.15)
    definitions.columns[1].width = Cm(9.15)
    for idx, (label, definition, color) in enumerate(legend_definitions):
        cell = definitions.cell(idx // 2, idx % 2)
        set_cell_margins(cell, top=35, start=75, bottom=35, end=75)
        set_cell_shading(cell, WHITE if idx // 2 % 2 == 0 else LIGHT_GREY)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        add_run(paragraph, f"{label}: ", bold=True, color=color, size=6.9)
        add_run(paragraph, definition, color=MID_GREY, size=6.9)
    for row in definitions.rows:
        prevent_row_split(row)
    set_table_borders(definitions, color="DDE3E8", size="3")

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    add_run(p, "How to read the risk signal: ", bold=True, color=TEAL, size=7.4)
    add_run(
        p,
        "The Overall bar shows 1,909 of 6,337 mature companies (30.1%) in the grouped Declining + inactive signal. This is not a separate account state: it combines Declining <€10 + Material Watch + At-risk + Churned proxy. The first two remain revenue-active; At-risk means 1–2 revenue-free months and Churned proxy means 3+ after prior revenue. It excludes Healthy/stable, Recovered, and Never monetized. This is a screening indicator—not confirmed churn. N is the mature-company population.",
        size=7.4,
        color=MID_GREY,
    )

    add_heading(document, "Alternatives considered—and rejected", level=2, space_before=1)
    rows = [
        ("Acquire more Business", "High revenue/company", "Small base; 73.3% funnel dropout; weak continuity; no CAC or KYB-reason data."),
        ("Broad funnel overhaul", "62.9% 30-day dropout", "Cannot distinguish rejection, delay, abandonment, or extraction effects; no channel cost."),
        ("Scale BTP immediately", "Strong scale and efficiency", "The data identifies the segment, but not the product or acquisition lever that caused its value."),
        ("Blanket churn campaign", "Retention feels urgent", "Decline and inactivity are different states; treatment should vary by state, exposure, and expected recovery."),
    ]
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Cm(4.0), Cm(4.1), Cm(10.2)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    headers = ["Alternative", "Why attractive", "Why it loses now"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=70, start=90, bottom=70, end=90)
        add_cell_text(cell, header, bold=True, size=7.8, color=WHITE, space_after=0)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(row_values):
            set_cell_margins(cells[idx], top=45, start=80, bottom=45, end=80)
            set_cell_shading(cells[idx], WHITE if len(table.rows) % 2 else LIGHT_GREY)
            add_cell_text(cells[idx], value, bold=(idx == 0), size=7.2, color=NAVY if idx == 0 else BLACK, space_after=0)
    set_table_borders(table, color="D8DEE3", size="3")

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Magnitude limit: ", bold=True, color=RED, size=7.7)
    add_run(
        p,
        "gross exposure can be sized; recoverability, uplift, contribution margin, treatment cost and causal impact cannot.",
        size=7.7,
        color=MID_GREY,
    )


def build_page_three(document: Document) -> None:
    add_page_title(
        document,
        "Page 3 • Prove and productionise",
        "Test the lever before scaling—and make the metrics trustworthy",
        "Association identifies where to test; randomisation establishes whether the lever works",
    )

    add_heading(document, "Lead experiment: revenue-deterioration protection", level=2, space_before=0)
    table = document.add_table(rows=5, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.0)
    table.columns[1].width = Cm(14.3)
    experiment = [
        ("Hypothesis", "A state-specific intervention after a material decline or first revenue-free month causes incremental recognised revenue versus business-as-usual."),
        ("Eligibility & design", "Freeze monthly eligibility; stratify by state, exposure/value, cohort and segment; randomise within strata. Prioritise Material Watch and first-gap At-risk: high-touch for the largest exposures, low-cost automation for the tail, and lower-touch monitoring for Churned proxy."),
        ("Primary metric", "Incremental mean total revenue per eligible company over 60–90 days. Report confidence intervals plus median and winsorised sensitivities because revenue is highly skewed."),
        ("Success", "Scale a treatment tier only if incremental contribution exceeds fully loaded intervention cost without guardrail deterioration. Estimate effects separately by state and value tier."),
        ("Guardrails", "Closures, complaints, support contacts, fraud/AML/risk outcomes, fairness, operational failures, and negative-fee adjustments."),
    ]
    for idx, (label, value) in enumerate(experiment):
        left, right = table.rows[idx].cells
        for cell in (left, right):
            set_cell_margins(cell, top=72, start=95, bottom=72, end=95)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(left, LIGHT_BLUE)
        set_cell_shading(right, WHITE if idx % 2 == 0 else LIGHT_GREY)
        add_cell_text(left, label, bold=True, color=TEAL, size=8.0, space_after=0)
        add_cell_text(right, value, size=7.8, space_after=0)
        prevent_row_split(table.rows[idx])
    set_table_borders(table, color="D8DEE3", size="3")

    add_heading(document, "Validate first", level=2, space_before=3)
    validation = document.add_table(rows=1, cols=3)
    validation.autofit = False
    for idx in range(3):
        validation.columns[idx].width = Cm(6.1)
    validations = [
        ("1", "Completeness", "Confirm missing row = zero; reconcile to Finance; keep May out until signup coverage is explained."),
        ("2", "Cause & actionability", "Add product/card usage, balances, current plan, status/closure, contact eligibility and cost; diagnose why revenue declined."),
        ("3", "KYB & privacy", "Confirm permitted persona use; minimise data; suppress small cells; never use this for adverse eligibility decisions."),
    ]
    for idx, (number, title, detail) in enumerate(validations):
        cell = validation.cell(0, idx)
        set_cell_shading(cell, LIGHT_GREY if idx != 2 else "FBEFEA")
        set_cell_margins(cell, top=90, start=105, bottom=90, end=105)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        add_run(p, number, bold=True, color=TEAL if idx != 2 else RED, size=12)
        add_run(p, f"  {title}", bold=True, color=NAVY, size=8.4)
        add_cell_text(cell, detail, size=7.5, color=BLACK, space_after=0)
    set_table_borders(validation, color=WHITE, size="7")

    add_heading(document, "Trusted production engine", level=2, space_before=3)
    flow = document.add_table(rows=1, cols=7)
    flow.autofit = False
    flow_items = [
        ("Raw", "snapshots"),
        ("→", ""),
        ("Typed", "company-month fact"),
        ("→", ""),
        ("Versioned", "metric marts"),
        ("→", ""),
        ("Governed", "semantic layer"),
    ]
    widths = [Cm(2.55), Cm(0.55), Cm(4.0), Cm(0.55), Cm(3.6), Cm(0.55), Cm(4.7)]
    for idx, width in enumerate(widths):
        flow.columns[idx].width = width
    for idx, (title, subtitle) in enumerate(flow_items):
        cell = flow.cell(0, idx)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=65, start=45, bottom=65, end=45)
        if title == "→":
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, title, bold=True, color=TEAL, size=12)
        else:
            set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            add_run(p, title, bold=True, color=NAVY, size=8.0)
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            add_run(p, subtitle, color=MID_GREY, size=6.9)
    set_table_borders(flow, color=WHITE, size="7")

    add_body(
        document,
        "Automated tests: unique company-month key; referential integrity; component-to-total and Finance reconciliation; freshness/completeness; date ordering; state-transition invariants; fixed cohort denominators; late-arriving-data/backfill controls.",
        size=7.75,
        space_after=1,
    )
    add_body(
        document,
        "Ownership: Finance—revenue; Product/Growth—lever and thresholds; Data—models/tests/lineage; Risk/Compliance/Privacy—KYB use; Operations/CRM—execution. Refresh after monthly close and version every definition change.",
        size=7.75,
        space_after=1,
    )

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "SQL appendix: ", bold=True, color=TEAL, size=7.7)
    add_run(p, "sql/revenue_concentration_analysis.sql • sql/presentation_kpis.sql • sql/account_health_playground.sql", size=7.7, color=MID_GREY)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Not concluded: ", bold=True, color=RED, size=7.7)
    add_run(p, "recoverable exposure, true churn, forecast LTV, causal plan effect, CAC/payback, profitability, or complete May performance.", size=7.7, color=MID_GREY)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "External sources: ", bold=True, color=TEAL, size=7.0)
    add_hyperlink(p, "Shine pricing", "https://www.shine.fr/tarifs/", size=7.0)
    add_run(p, " • ", size=7.0, color=MID_GREY)
    add_hyperlink(p, "Overdraft policy", "https://help.shine.fr/fr/articles/1179200-compte-a-decouvert", size=7.0)
    add_run(p, " • ", size=7.0, color=MID_GREY)
    add_hyperlink(p, "E-invoice reform", "https://www.economie.gouv.fr/tout-savoir-sur-la-facturation-electronique-pour-les-entreprises", size=7.0)
    add_run(p, " • ", size=7.0, color=MID_GREY)
    add_run(p, "full research note: docs/research/SHINE_MARKET_RESEARCH_AND_STRATEGY.md", size=7.0, color=MID_GREY)


def build_document() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    ranking_chart = ASSET_DIR / "segment_opportunity_ranking.png"
    health_chart = ASSET_DIR / "segment_health_stacked.png"
    make_segment_ranking_chart(ranking_chart)
    make_health_chart(health_chart)

    document = Document()
    configure_document(document)
    build_page_one(document, ranking_chart)
    document.add_page_break()
    build_page_two(document, health_chart)
    document.add_page_break()
    build_page_three(document)
    document.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    output = build_document()
    print(output)
